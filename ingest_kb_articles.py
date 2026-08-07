####################################################################################################
# Project name      : Outlook Support Classification Agent -- Foundry v2 (ServiceNow-shaped KB)    #
# Business owner    : <fill: business owner / team>                                                #
# Notebook Author   : <fill: author name / team>                                                   #
# Date              : <fill: date>                                                                 #
#                                                                                                  #
# Purpose of file:                                                                                 #
#   1. ONE-TIME LOCAL ingest: reads one .docx per KB article and builds kb_index.json.             #
#   2. KB number comes from the filename (e.g. KB0024755.docx); full text -> Article Content.      #
#   3. Output mirrors the ServiceNow search-result shape (POC == prod agent/prompt).               #
#                                                                                                  #
#   Run locally def:- python ingest_kb_articles.py in the POC environment; NOT deployed.           #
#                                                                                                  #
# Source:-                                                                                         #
#   - python-docx (docx) reads the .docx text; imported lazily with a clear error if missing.      #
#   - json / pathlib / sys are the Python standard library.                                        #
####################################################################################################

# ============================================ Imports =============================================
from __future__ import annotations  # Enable postponed evaluation of type annotations (PEP 563)   # future import

import json  # Serialise the built index to kb_index.json                                          # stdlib json
import sys  # Exit with a clear status on setup problems                                           # stdlib sys
from pathlib import Path  # Filesystem paths resolved against this script's folder                 # stdlib path
from typing import Any, Optional  # Generic type hints for the record dicts                        # typing helpers

try:  # python-docx extracts the article text from each .docx                                      # optional import
    import docx  # python-docx                                                                     # docx lib
except ImportError:  # Fail with clear guidance if it is not installed                              # missing dep
    docx = None  # Sentinel checked in main()                                                      # sentinel

# ========================================= Configuration =========================================
_HERE = Path(__file__).resolve().parent  # This script's folder (the deployment folder)            # base folder
_SOURCE_DIR = _HERE / "kb_source"  # Folder holding one <KBNUMBER>.docx per article                # source dir
_OUTPUT_PATH = _HERE / "kb_index.json"  # The ServiceNow-shaped index the service loads at runtime  # output file
_MAX_SNIPPET_CHARS = 240  # Cap for the top-level "text" preview snippet                            # snippet cap


# =========================================== Helpers =============================================
def _read_docx_text(docx_path: Path) -> str:  # Extract paragraph + table text from one .docx
    """Extract the full readable text (paragraphs + tables) from a .docx, verbatim.

    What this function is:
        - The deterministic reader: it returns the article's text exactly as written, so the
          "Article Content" in the index is the real document (no summarisation, no LLM).

    Why this exists:
        - The agent grounds its choice on the FULL article content; fidelity matters, so we read
          the document rather than paraphrasing it.

    Security and production notes:
        1. Read-only text extraction; images/embedded objects are ignored (not needed here).

    Args:
        docx_path: Path to the .docx file.

    Returns:
        The document text (paragraphs, then table rows), newline-joined.

    Example:
        >>> _read_docx_text(Path("kb_source/KB0024755.docx"))  # doctest: +SKIP
        'Shared Mailbox Not Working ...'
    """
    document = docx.Document(str(docx_path))  # Open the .docx                                       # open docx
    lines: list[str] = [p.text for p in document.paragraphs if p.text and p.text.strip()]  # Non-empty paragraphs  # paragraphs
    for table in document.tables:  # Include any table cell text                                    # tables
        for row in table.rows:  # Each row                                                          # per row
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]  # Non-empty cells  # cells
            if cells:  # Only rows with content                                                     # has content?
                lines.append(" | ".join(cells))  # Join the row's cells                             # join row
    return "\n".join(lines)  # The article text                                                     # joined text


def _build_record(kb_number: str, title: str, content: str) -> dict[str, Any]:  # One ServiceNow-shaped record
    """Assemble one ServiceNow-shaped KB record (columns list) for a single article.

    What this function is:
        - The record builder: it mirrors the ServiceNow search-result entry shape so the agent
          sees the SAME structure in the POC (local file) as in production (the KB search API).

    Why this exists:
        - Keeping the shape identical means only the SOURCE changes when moving to production
          (local file -> API); the agent, prompt, and grounding stay the same.

    Security and production notes:
        1. `number` and `Article Content` are the load-bearing fields; sys_id / dates / category are
           placeholders here (the real API fills them) and are not relied on by the agent.

    Args:
        kb_number: The KB number (from the filename), e.g. "KB0024755".
        title: A human title / short description for the article.
        content: The FULL article text (verbatim).

    Returns:
        A dict shaped like one ServiceNow searchResults entry.

    Example:
        >>> _build_record("KB0024755", "Shared Mailbox ...", "Shared Mailbox ...")  # doctest: +SKIP
        {'sysId': '', 'table': 'kb_knowledge', ...}
    """
    snippet = content[:_MAX_SNIPPET_CHARS].strip()  # A short preview for the top-level "text" field  # snippet
    return {  # The record, matching the ServiceNow searchResults entry shape                       # record
        "sysId": "",  # Placeholder (the real ServiceNow API supplies this)                         # sys id
        "table": "kb_knowledge",  # Constant table name, as ServiceNow returns                      # table
        "text": snippet,  # Highlighted-snippet stand-in (preview only)                             # snippet field
        "title": title,  # The article title / short description                                    # title
        "columns": [  # The column entries the agent reads (number + Article Content are key)       # columns
            {"fieldName": "short_description", "label": "Short description|Purpose", "value": title, "displayValue": title},  # short desc
            {"fieldName": "number", "label": "Number", "value": kb_number, "displayValue": kb_number},  # KB number (the id)
            {"fieldName": "sys_id", "label": "sys_id", "value": "", "displayValue": ""},  # placeholder sys_id
            {"fieldName": "text", "label": "Article Content", "value": content, "displayValue": content},  # FULL content
            {"fieldName": "sys_updated_on", "label": "Updated", "value": "", "displayValue": ""},  # placeholder date
            {"fieldName": "kb_category", "label": "Category", "value": "", "displayValue": ""},  # placeholder category
        ],
    }


def _title_for(kb_number: str, content: str) -> str:  # Derive a title / short description
    """Return the article's first non-empty line as its title, or the KB number as a fallback."""
    for line in content.splitlines():  # Walk the lines top-down                                     # per line
        if line.strip():  # First non-empty line                                                    # non-empty?
            return line.strip()[:180]  # Use it (capped) as the title                               # title
    return kb_number  # No text -> fall back to the number                                          # fallback


# =========================================== Entry point =========================================
def main() -> int:  # Build kb_index.json from the .docx files in ./kb_source
    """Read every <KBNUMBER>.docx in ./kb_source and write kb_index.json (ServiceNow shape).

    What this function is:
        - The one-time local entry point. It scans ./kb_source, reads each article verbatim, and
          writes the ServiceNow-shaped index the runtime KB source loads.

    Why this exists:
        - To produce the exact JSON structure the agent expects, from your own KB documents, with
          NO LLM and NO external dependency (run it once in VS Code in the POC environment).

    Security and production notes:
        1. Local, read-only ingest; content is verbatim (no invention). Not deployed with the app
           (excluded via .funcignore).

    Args:
        None.

    Returns:
        Process exit code (0 on success, non-zero on a setup problem).

    Example:
        >>> main()  # doctest: +SKIP
        0
    """
    if docx is None:  # python-docx is required to read the .docx files                             # dep guard
        print("ERROR: python-docx is not installed. Run: pip install python-docx", file=sys.stderr)  # remedy
        return 2  # Non-zero exit                                                                   # exit
    if not _SOURCE_DIR.is_dir():  # The source folder must exist                                     # source guard
        print(f"ERROR: source folder not found: {_SOURCE_DIR} (create it and add <KBNUMBER>.docx files).", file=sys.stderr)  # msg
        return 2  # Non-zero exit                                                                    # exit

    records: list[dict[str, Any]] = []  # Collected ServiceNow-shaped records                        # records
    for docx_path in sorted(_SOURCE_DIR.glob("*.docx")):  # Every article file (sorted for stable output)  # per file
        if docx_path.name.startswith("~$"):  # Skip Word's temporary lock files                     # temp file?
            continue  # Ignore it                                                                   # skip
        kb_number = docx_path.stem.strip()  # KB number = the filename without extension (e.g. KB0024755)  # number
        content = _read_docx_text(docx_path)  # The FULL article text, verbatim                     # content
        if not content.strip():  # An empty document is not usable                                  # empty?
            print(f"  ! {docx_path.name}: no readable text -- skipping")  # note the skip           # skip note
            continue  # Move on                                                                     # next
        title = _title_for(kb_number, content)  # Derive the title / short description               # title
        records.append(_build_record(kb_number, title, content))  # Add the record                  # add
        print(f"  + {kb_number}: {len(content)} chars")  # Progress line                            # progress

    if not records:  # Nothing was ingested                                                         # empty result?
        print(f"No <KBNUMBER>.docx files found in {_SOURCE_DIR}.")  # note                          # note
        return 1  # Non-zero exit                                                                   # exit

    _OUTPUT_PATH.write_text(json.dumps({"searchResults": records}, indent=2, ensure_ascii=False), encoding="utf-8")  # Write index
    print(f"\nWrote {len(records)} article(s) to {_OUTPUT_PATH}")  # Summary                        # summary
    print("Review kb_index.json (number + Article Content are the load-bearing fields), then deploy the folder.")  # hint
    return 0  # Success                                                                             # exit


if __name__ == "__main__":  # Allow `python ingest_kb_articles.py`                                  # script guard
    raise SystemExit(main())  # Exit with the return code                                           # run
